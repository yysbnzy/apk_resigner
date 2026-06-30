#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 APK 签名替换工具 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcPr.append(element)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), 'auto')

def create_readme_docx():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('APK 签名替换工具 - 使用说明文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('版本: v2.3.0 | 更新日期: 2026-06-25')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # 概述
    doc.add_heading('一、工具概述', 1)
    doc.add_paragraph(
        'APK 签名替换工具是一款用于 Android 应用签名验证测试的便携工具，'
        '支持 ADB 设备管理、应用扫描、备份还原、一键签名替换等功能。'
    )
    
    # 特点
    doc.add_heading('二、主要特点', 1)
    features = [
        '零依赖运行：无需安装 Android SDK、JDK，所有工具内置（可选纯 Python 模式）',
        '单文件 EXE：复制到任何 Windows 电脑都能直接运行',
        'GUI 界面：可视化操作，无需命令行',
        'ADB 集成：设备连接、应用扫描、备份还原、一键处理',
        '自动检测签名：选择 APK 后自动识别 V1/V2/V3/V4 签名方案',
        '备份还原：导出原版 APK 备份，随时还原'
    ]
    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    # 文件说明
    doc.add_heading('三、文件说明', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '文件'
    hdr_cells[1].text = '用途'
    
    files = [
        ('apk_resigner_gui.py', 'GUI 主程序（含 ADB 模块集成）'),
        ('apk_resigner.py', '命令行版本'),
        ('quick_sign_replace.py', '快速签名替换（不解包）'),
        ('pure_python_sign.py', '纯 Python 签名引擎（无需 JDK）'),
        ('adb_manager.py', 'ADB 设备管理模块'),
        ('backup_manager.py', '备份/还原管理模块'),
        ('install_manager.py', 'APK 安装策略模块'),
        ('build_portable.py', '自动收集工具并打包为 EXE'),
        ('build_portable.bat', 'Windows 一键打包批处理'),
        ('_tools/', '依赖工具目录（打包时自动包含）'),
    ]
    for fname, desc in files:
        row_cells = table.add_row().cells
        row_cells[0].text = fname
        row_cells[1].text = desc
    
    doc.add_page_break()
    
    # 使用场景
    doc.add_heading('四、使用场景', 1)
    
    doc.add_heading('场景 1：不连接设备，本地 APK 重签名', 2)
    steps1 = [
        '打开工具，选择本地 APK 文件',
        '工具自动检测签名方案（V1/V2/V3/V4）',
        '选择操作：',
        '  • 修改内容+签名：反编译 → 修改 → 重打包 → zipalign → 签名',
        '  • 快速签名替换：去除原签名 → 重新签名（不改内容）',
        '查看签名结果和验证报告'
    ]
    for step in steps1:
        p = doc.add_paragraph(step, style='List Number' if step[0].isdigit() else 'List Bullet')
    
    doc.add_heading('场景 2：连接车机，一键导出→备份→签名→安装测试', 2)
    steps2 = [
        '连接 Android 设备（USB 或网络 ADB）',
        '切换到「📱 ADB设备」标签页，点击「刷新」自动检测设备',
        '切换到「📦 应用列表」标签页，点击「扫描」获取应用列表',
        '选择目标应用，点击「一键处理」',
        '自动执行：导出 → 备份 → 签名',
        '签名完成后，使用「覆盖安装」或「卸载后安装」测试',
        '预期结果：安装失败（签名不匹配），验证签名校验机制工作正常'
    ]
    for step in steps2:
        p = doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('场景 3：备份还原，恢复原始 APK', 2)
    steps3 = [
        '切换到「💾 备份还原」标签页',
        '选择应用，查看已有备份列表',
        '选中备份，点击「还原选中」',
        '确认还原信息后，工具直接安装原始 APK（无需签名）',
        '恢复完成'
    ]
    for step in steps3:
        p = doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('⚠️ 注意：还原备份不需要签名！')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph('还原时使用的是原版 APK（带原始厂商签名），直接安装即可。')
    
    doc.add_page_break()
    
    # 功能详解
    doc.add_heading('五、功能详解', 1)
    
    doc.add_heading('5.1 本地 APK 签名', 2)
    doc.add_paragraph('不连接设备时，直接使用主界面选择的 APK 文件进行重签名。')
    
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Light Grid Accent 1'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = '按钮'
    hdr_cells[1].text = '功能说明'
    
    buttons = [
        ('🔧 修改内容+签名', '反编译 APK，修改内容后重新打包并签名'),
        ('⚡ 快速签名替换', '不解包，直接去除原签名并重新签名'),
        ('🔍 验证签名', '检查 APK 签名状态和对齐情况'),
        ('📊 签名对比', '对比两个 APK 的签名差异'),
    ]
    for btn, desc in buttons:
        row_cells = table2.add_row().cells
        row_cells[0].text = btn
        row_cells[1].text = desc
    
    doc.add_heading('5.2 ADB 设备管理', 2)
    
    doc.add_heading('📱 设备连接', 3)
    doc.add_paragraph('自动检测并连接 USB 设备，显示设备信息（型号、Android 版本、序列号等）。')
    
    doc.add_heading('📦 应用列表', 3)
    doc.add_paragraph('扫描设备上的所有应用（第三方/系统），支持搜索过滤，查看应用详情。')
    
    doc.add_heading('💾 备份还原', 3)
    p = doc.add_paragraph()
    p.add_run('创建备份：').bold = True
    p.add_run('导出 APK 并保存到备份目录')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('还原备份：').bold = True
    p.add_run('直接安装原版 APK（无需签名）')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('清理旧备份：').bold = True
    p.add_run('自动清理过期备份')
    
    doc.add_heading('📝 ADB 日志', 3)
    doc.add_paragraph('左右分区设计：')
    doc.add_paragraph('左侧：操作日志 - 显示工具操作记录和状态', style='List Bullet')
    doc.add_paragraph('右侧：ADB 命令 - 显示 ADB 命令输入输出，便于调试', style='List Bullet')
    
    doc.add_page_break()
    
    # 签名方案说明
    doc.add_heading('六、签名方案说明', 1)
    
    table3 = doc.add_table(rows=1, cols=4)
    table3.style = 'Light Grid Accent 1'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '方案'
    hdr_cells[1].text = 'Android 版本'
    hdr_cells[2].text = '说明'
    hdr_cells[3].text = '签名文件'
    
    schemes = [
        ('V1 (JAR)', '5.0-6.0', '基于 JAR 签名', 'META-INF/CERT.*'),
        ('V2', '7.0+', 'APK Signing Block，整文件签名', 'APK 文件末尾'),
        ('V3', '9.0+', 'V2 + 证书轮换支持', 'APK 文件末尾'),
        ('V4', '11.0+', '增量签名，用于 APEX', 'APK 文件末尾'),
    ]
    for scheme, ver, desc, loc in schemes:
        row_cells = table3.add_row().cells
        row_cells[0].text = scheme
        row_cells[1].text = ver
        row_cells[2].text = desc
        row_cells[3].text = loc
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('工具会自动检测原 APK 使用的签名方案，并保持一致。')
    run.italic = True
    
    doc.add_page_break()
    
    # 常见问题
    doc.add_heading('七、常见问题解答', 1)
    
    faqs = [
        ('Q: 还原备份需要签名吗？',
         '不需要。还原备份直接安装原版 APK（带原始厂商签名），无需重新签名。'),
        
        ('Q: 一键处理报 INSTALL_PARSE_FAILED_NO_CERTIFICATES？',
         '此错误已在 v2.2.9 修复。旧版本使用 zipfile 追加模式破坏 V2 签名块，'
         '新版本改为：解压 → 清除 META-INF → 修改 → 重打包 → zipalign → 签名。'),
        
        ('Q: 签名后的 APK 安装失败是正常的吗？',
         '是的！签名后的 APK 使用测试密钥签名，与原始厂商签名不匹配。'
         '系统拒绝安装是预期行为，说明签名校验机制工作正常。'),
        
        ('Q: ADB 连接不上设备怎么办？',
         '1. 检查 USB 线是否连接正常\\n'
         '2. 在设备上开启「开发者选项」和「USB 调试」\\n'
         '3. 允许电脑调试授权（设备上会弹出提示）\\n'
         '4. 尝试重新插拔 USB 线\\n'
         '5. 检查设备驱动是否安装'),
        
        ('Q: 打包后的 EXE 有多大？',
         '约 50-150MB，取决于内置的 JDK 大小。'),
        
        ('Q: 可以去掉 Java 依赖吗？',
         '可以！启用纯 Python 模式：无需 JDK/Android SDK，仅支持 V1 (JAR) 签名，需要 pip install cryptography。'),
        
        ('Q: 为什么启动比较慢？',
         '单文件 EXE 启动时需要解压内置的 _tools/ 到临时目录，首次启动约 5-10 秒。'),
        
        ('Q: 杀毒软件报毒？',
         'PyInstaller 打包的 EXE 常被误报，添加信任即可。'),
    ]
    
    for question, answer in faqs:
        p = doc.add_paragraph()
        run = p.add_run(question)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)
        doc.add_paragraph(answer)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # 更新日志
    doc.add_heading('八、版本更新日志', 1)
    
    versions = [
        ('v2.3.0 (2026-06-25)', [
            '修复备份布局显示问题',
            'ADB 日志分割为操作日志和命令日志',
            '添加命令执行记录器',
        ]),
        ('v2.2.9 (2026-06-26)', [
            '重要修复：签名流程重构',
            '_full_process 改为：解压 → 清除 META-INF → 添加 test.txt → 重打包',
            '_sign_with_scheme 签名前自动清理旧签名残留',
            '避免 zipfile "a" 模式破坏 V2 签名块位置',
        ]),
        ('v2.2.1 (2026-06-18)', [
            '添加 V1-only 签名支持',
            '纯 Python 签名引擎（无需 JDK）',
            '自动检测签名方案',
            '5 按钮布局：修改+签名 / V1 / V2 / V2+V3 / V4',
        ]),
        ('v2.0.0 (2026-06-18)', [
            'ADB 扩展功能 v2.0',
            '新增模块：ADBManager / BackupManager / InstallManager',
            '4 个 Notebook 标签页：设备/应用/备份/日志',
            '一键流程：导出→备份→签名→安装测试',
        ]),
        ('v1.2.0 (2026-06-17)', [
            '签名对比工具',
            '使用说明弹窗',
            'GitHub Release 发布',
        ]),
        ('v1.1 (2026-06-17)', [
            '便携版支持',
            '内置工具路径',
            '单文件 EXE 打包',
        ]),
        ('v1.0 (2026-06-17)', [
            '基础命令行工具',
            'GUI 界面',
            '签名替换功能',
        ]),
    ]
    
    for version, changes in versions:
        doc.add_heading(version, 2)
        for change in changes:
            doc.add_paragraph(change, style='List Bullet')
    
    doc.add_page_break()
    
    # 技术细节
    doc.add_heading('九、技术细节', 1)
    
    doc.add_heading('9.1 内置工具路径检测优先级', 2)
    doc.add_paragraph('1. PyInstaller 临时目录（sys._MEIPASS/_tools/）← 单文件 EXE 模式', style='List Number')
    doc.add_paragraph('2. EXE 同级目录（exe_dir/_tools/）← 目录模式', style='List Number')
    doc.add_paragraph('3. 源码目录（script_dir/_tools/）← 源码运行', style='List Number')
    doc.add_paragraph('4. 系统 PATH ← 回退', style='List Number')
    
    doc.add_heading('9.2 一键处理流程', 2)
    doc.add_paragraph('1. 导出：从设备 adb pull 原版 APK', style='List Number')
    doc.add_paragraph('2. 备份：保存原版 APK 到备份目录', style='List Number')
    doc.add_paragraph('3. 签名：修改内容并替换签名', style='List Number')
    doc.add_paragraph('4. 安装测试：验证签名校验机制', style='List Number')
    
    p = doc.add_paragraph()
    run = p.add_run('注意：第 4 步预期安装失败，用于验证系统签名校验机制。')
    run.italic = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_heading('9.3 最小化 JRE', 2)
    doc.add_paragraph('脚本从本地 JDK 复制以下文件到 _tools/java/：')
    doc.add_paragraph('bin/java.exe, keytool.exe, jarsigner.exe', style='List Bullet')
    doc.add_paragraph('lib/ 目录（运行时库，约 50-100MB）', style='List Bullet')
    doc.add_paragraph('conf/ 目录（配置文件）', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('删除了 demo/, sample/, man/, src.zip 等不必要文件以减小体积。')
    
    doc.add_page_break()
    
    # 安全提示
    doc.add_heading('十、安全提示', 1)
    
    p = doc.add_paragraph()
    run = p.add_run('⚠️ 本工具仅用于测试 Android 应用完整性校验，请勿用于非法用途。')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    doc.add_paragraph('替换签名后的 APK：')
    doc.add_paragraph('无法通过原开发者签名校验', style='List Bullet')
    doc.add_paragraph('无法通过 Google Play Protect / SafetyNet', style='List Bullet')
    doc.add_paragraph('系统级应用通常无法安装', style='List Bullet')
    
    # 联系信息
    doc.add_page_break()
    doc.add_heading('十一、联系信息', 1)
    
    doc.add_paragraph('GitHub: https://github.com/yysbnzy')
    doc.add_paragraph('项目地址: https://github.com/yysbnzy/apk_resigner')
    doc.add_paragraph('Release 页面: https://github.com/yysbnzy/apk_resigner/releases')
    
    # 保存
    output_path = r'C:\Users\Administrator\apk_resigner-main\APK签名替换工具-使用说明.docx'
    doc.save(output_path)
    print(f'Word 文档已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    create_readme_docx()
